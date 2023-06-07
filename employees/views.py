from django.shortcuts import render
from django.template import loader
from django.http import HttpResponse
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from django.http import FileResponse
import openai
import re
import requests
import os
import docx
from bs4 import BeautifulSoup
from docx import Document
from django.http import Http404

# Create your views here.
def main(request):
   template = loader.get_template('index.html')
   return HttpResponse(template.render())
# def jobseeker(request):
#    template = loader.get_template('jobseeker.html')
#    return HttpResponse(template.render(request))
def jobseeker(request):
    return render(request,'jobseeker.html')
def resume(request):
    return render(request,'resumefinder.html')

# def scrape_jobs():
#     url = 'https://www.freshersworld.com/jobs/jobsearch/python-jobs-in-hyderabad?experience=12'
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, 'html.parser')

#     job_listings = []

#     # Find the HTML elements containing job information
#     job_elements = soup.find_all('div', class_='job-container')
#     print(job_elements)

#     for job_element in job_elements:
#         # Extract job details
#         job_title = job_elements.find('span', class_='wrap-title').text
#         company_name = job_elements.find('h3', class_='latest-jobs-title font-16 margin-none inline-block company-name').text
#         job_location = job_elements.find('span', class_='job-location display-block modal-open job-details-span').text
#         qualifications = job_elements.find('span', class_='qualifications display-block modal-open pull-left job-details-span').text
                    
#         apply = job_element.find('a', class_='apply-now').text.strip()
#         link = job_element.find('a', class_='list-job-link')['href']

#         # Add job details to the list
#         job_listings.append({
#             'title': job_title,
#             'location': job_location,
#             'qualification': qualifications,
#             'apply': apply,
#             'link': link
#         })
#         print(job_listings)

#     return job_listings
    
# def job_search(request):
#     if request.method == 'POST':
#          description = request.POST.get('description', '')
#          print("Search Description:", description)
#         # Call the scrape_jobs() function to get the job listings
#          job_listings = scrape_jobs()

#          return render(request, 'jobseeker.html', {'job_listings': job_listings,'description':description})

#     return render(request, 'jobseeker.html')
def get_resume_names(folder_path):
    resume_names = [doc for doc in os.listdir(folder_path) if doc.endswith(('.txt', '.pdf', '.docx')) and not doc.startswith('~$')]
    return resume_names
openai.api_key = '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'

def filtered_files(request,technologies, experience):
    
    if request.method == 'POST':
        technologies = request.POST.get('technologies')
        experience = request.POST.get('experience')
    
    
 
        description=' '
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=f"Summarize  the enter description in a short form and get the technologies:\n{description}\n",
            temperature=0.5,
            max_tokens=100,
            n=1,
            stop=None,
            timeout=10,
        )

        summary = response.choices[0].text.strip()
        lemmatizer = WordNetLemmatizer()

        lemmatized_words = [lemmatizer.lemmatize(word) for word in summary]
        #  print(lemmatized_words)
        #  print(summary)
        stop_words =stopwords.words("english")
        stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
        folder_path = "C:/Users/singl/OneDrive/Desktop/docs"
        
        
        matching_files = []
        matching_files1 = []
        matching_files2 = []
        files = os.listdir(folder_path)
        for file in files:
            if file.endswith('.docx'):
                document = Document(os.path.join(folder_path, file))

                text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                text = text.lower()

            
                pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
                
                matches = re.findall(pattern, text)
                
                numeric_matches = [match[0] for match in matches if match[0]]
                word_matches = [match[1] for match in matches if match[1]]
                
                totals = numeric_matches + word_matches

                print(totals)


                summary_pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"

                matches = re.findall(summary_pattern, description)
                
                
                numeric_matches1 = [match[0] for match in matches if match[0]]
                word_matches1 = [match[1] for match in matches if match[1]]
                
                sum_totals = numeric_matches1 + word_matches1

                print(sum_totals)


                state = ""
                for sum_total in sum_totals:
                    for total in totals:
                        if sum_total in total:
                            state = str("Experience {} years match " .format(total))
                            print(state)
                    

                tokens = word_tokenize(summary.lower())
                filtered_words = []
                for token in tokens:
                    if token not in stop_words:
                        filtered_words.append(token)
                print(filtered_words)
                num_keywords = len(filtered_words)
                matching_keywords = []
                for word in filtered_words:
                    if word in text:
                        matching_keywords.append(word)



                matching_keywords1 = len(matching_keywords)
                print(matching_keywords)
                matching_percentage = matching_keywords1 / num_keywords * 100
                print(matching_percentage)
                matching = []

                if matching_percentage>=1:

                    if matching_percentage >= 75:
                        matching_files.append((file, matching_percentage,matching_keywords,state))
                    elif matching_percentage>=50 and matching_percentage<75:
                        matching_files1.append((file, matching_percentage,matching_keywords,state))
                    elif matching_percentage>=1 and matching_percentage<50:
                    # else:
                        matching_files2.append((file, matching_percentage,matching_keywords,state)) 
                    else:
                        print("No Files Found")

                else:
                    print("No Files Found")
        
        sorted_files1 = sorted(matching_files, key=lambda x: x[1], reverse=True)
    
        sorted_files2 = sorted(matching_files1, key=lambda x: x[1], reverse=True)
    
        sorted_files3 = sorted(matching_files2, key=lambda x: x[1], reverse=True)
        
        filtered_files=[]
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

        for file_info in sorted_files1:
            file_name = file_info[0]
            file_path = folder_path + file_name
            filtered_files.append(file_path)

        print("filtered_files with sorted_files1 start")
        print(filtered_files)
        print("filtered_files with sorted_files1 end")




        filtered_files1=[]
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

        for file_info in sorted_files2:
            file_name = file_info[0]
            file_path = folder_path + file_name
            filtered_files1.append(file_path)

        print("filtered_files1 with sorted_files2 start")
        print(filtered_files1)
        print("filtered_files1 with sorted_files2 end")

        filtered_files2=[]
        folder_path = "C:/Users/singl/OneDrive/Desktop/docs"

        for file_info in sorted_files3:
            file_name = file_info[0]
            file_path = folder_path + file_name
            filtered_files2.append(file_path)
        filtered_files = [os.path.basename(file) for file in filtered_files]
        filtered_files1 = [os.path.basename(file) for file in filtered_files1]
        filtered_files2 = [os.path.basename(file) for file in filtered_files2]
        # filtered_files = [file for file in filtered_files]
        # filtered_files1 = [file for file in filtered_files1]
        # filtered_files2 = [file for file in filtered_files2]
        # filtered_files = [os.path.relpath(file, settings.MEDIA_ROOT) for file in filtered_files]

        print("filtered_files2 with sorted_files3 start")
        print(filtered_files2)
        print("filtered_files2 with sorted_files3 start")
        

        matching_files_str = "\n".join([f"{file}         {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state} " for file, percentage,matching_keywords,state in sorted_files1])
        matching_files_str1 = "\n".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files2])
        matching_files_str2 = "\n".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files3])
        

        count = len(matching_files + matching_files1 + matching_files2)

        print(f" {count} files were filtered ".format(count))
        
        context = {
            'count': count,
            'filtered_files': filtered_files,
            'matching_files_str': matching_files_str,
            'filtered_files1': filtered_files1,
            'matching_files_Str1':matching_files_str1,
            'filtered_files2': filtered_files2,
            'matching_files_str2':matching_files_str2,
        }
        print(f"{count} files were filtered")
    # return count,filtered_files,matching_files_str,filtered_files1,matching_files_str1,filtered_files2,matching_files_str2

    return render(request, 'resumefinder.html', context,experience,technologies)
  
        
    
# from django.http import HttpResponse
# import os
# from django.http import HttpResponse, FileResponse
# from django.conf import settings
# def download_file(request, filename):
#     file_path = os.path.join('C:/Users/sp13/OneDrive/Desktop/Resumes/', filename)  # Replace 'path_to_folder' with the actual folder path
#     if os.path.exists(file_path):
#         with open(file_path, 'rb') as file:
#             response = HttpResponse(file.read(), content_type='application/docx')
#             response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
#             return response
#     else:
#         raise Http404("File does not exist")




    
  #return count,filtered_files,matching_files_str,filtered_files1,matching_files_str1,filtered_files2,matching_files_str2
# def download_file(request, filename):
#     file_path = f"C:/Users/sp13/OneDrive/Desktop/Resumes/{filename}"  # Update the file path based on your actual directory

#     try:
#         with open(file_path, 'rb') as file:
#             response = FileResponse(file)
#             response['Content-Disposition'] = f'attachment; filename="{filename}"'
#             return response
#     except FileNotFoundError:
#         raise Http404("File does not exist")     

# from django.http import FileResponse, Http404
# import os

# def download_file(request, filename):
#     file_path = os.path.join('C:/Users/singl/OneDrive/Desktop/Resumes/', filename)  # Update with the actual folder path

#     if os.path.exists(file_path):
#         with open(file_path, 'rb') as file:
#             response = FileResponse(file)
#             response['Content-Disposition'] = f'attachment; filename="{filename}"'
#             return response

#     return render(request, 'resumefinder.html')
import mimetypes
import zipfile
def download_file(request, filename):
    folder_path = 'C:/Users/singl/OneDrive/Desktop/docs/'  # Replace with the actual folder path
    file_path = os.path.join(folder_path, filename)

    if os.path.isfile(file_path):
        fl = open(file_path, 'rb')
        mime_type, _ = mimetypes.guess_type(file_path)
        response = HttpResponse(fl, content_type=mime_type)
        response['Content-Disposition'] = f"attachment; filename={filename}"
        return response
    else:
        raise Http404("File does not exist")

# def download_files(request):
#     folder_path = 'C:/Users/singl/OneDrive/Desktop/Resumes/'  # Replace with the actual folder path
#     files = request.POST.getlist('files')  # Assuming you have a list of selected files from the form

#     zip_path = os.path.join(folder_path, 'download.zip')
#     zip_file = zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED)

#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         if os.path.isfile(file_path):
#             zip_file.write(file_path, os.path.basename(file_path))

#     zip_file.close()

#     if os.path.isfile(zip_path):
#         fl = open(zip_path, 'rb')
#         mime_type, _ = mimetypes.guess_type(zip_path)
#         response = HttpResponse(fl, content_type=mime_type)
#         response['Content-Disposition'] = 'attachment; filename="download.zip"'
#         return response
#     else:
#         raise Http404("No files selected or does not exist")

# def search1(request):
#     if request.method == 'POST':
#         description = request.POST.get('description1')
#         def filtered_files(request):
 
        
#             response = openai.Completion.create(
#                 engine="text-davinci-003",
#                 prompt=f"Summarize  the enter description in a short form and get the technologies:\n{description}\n",
#                 temperature=0.5,
#                 max_tokens=100,
#                 n=1,
#                 stop=None,
#                 timeout=10,
#             )

#             summary = response.choices[0].text.strip()
#             lemmatizer = WordNetLemmatizer()

#             lemmatized_words = [lemmatizer.lemmatize(word) for word in summary]
#             #  print(lemmatized_words)
#             #  print(summary)
#             stop_words =stopwords.words("english")
#             stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
#             folder_path = "C:/Users/singl/OneDrive/Desktop/docs"
            
            
#             matching_files = []
#             matching_files1 = []
#             matching_files2 = []
#             files = os.listdir(folder_path)
#             for file in files:
#                 if file.endswith('.docx'):
#                     document = Document(os.path.join(folder_path, file))

#                     text = "\n".join([paragraph.text for paragraph in document.paragraphs])
#                     text = text.lower()
            
#         return render(request, 'results.html', {'description': description, })
#     return render(request,'resumefinder.html')
# def searches(technologies, experience):
#     lemmatizer = WordNetLemmatizer()
#     folder_path = "C:/Users/sp13/OneDrive/Desktop/docs"

#     matching_files = []
    
#     files = os.listdir(folder_path)
#     for file in files:
#         if file.endswith('.docx'):
#             document = Document(os.path.join(folder_path, file))

#             text = "\n".join([paragraph.text for paragraph in document.paragraphs])
#             text = text.lower()

#             # Perform filtering based on technologies
#             technologies_list = technologies.split(',')
#             for tech in technologies_list:
#                 if tech.strip().lower() in text:
#                     matching_files.append(file)
#                     break

#             # Perform filtering based on experience
#             pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
#             matches = re.findall(pattern, text)
#             numeric_matches = [match[0] for match in matches if match[0]]
#             word_matches = [match[1] for match in matches if match[1]]
#             totals = numeric_matches + word_matches

#             for total in totals:
#                 if experience.strip().lower() in total.lower():
#                     matching_files.append(file)
#                     break

#     return matching_files

def search_resumes(request):
    if request.method == 'GET':
        technologies = request.GET.get('technologies', '')
        experience = request.GET.get('experience', '')
        param1 = request.GET.get('param1', '')
        param2 = request.GET.get('param2', '')

        count, filtered_files, matching_files_str, filtered_files1, matching_files_str1, filtered_files2, matching_files_str2 = filtered_files(technologies, experience, param1, param2)

        context = {
            'count': count,
            'filtered_files': filtered_files,
            'matching_files_str': matching_files_str,
            'filtered_files1': filtered_files1,
            'matching_files_str1': matching_files_str1,
            'filtered_files2': filtered_files2,
            'matching_files_str2': matching_files_str2,
            'technologies': technologies,
            'experience': experience,
            'param1': param1,
            'param2': param2
        }

        return render(request, 'search_results.html', context)
            
    




